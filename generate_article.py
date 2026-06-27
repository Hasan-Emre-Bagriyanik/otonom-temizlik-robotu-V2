"""Otonom temizlik robotu projesinin makalesini Word dosyasi olarak uretir."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            el = tcBorders.find(qn(f'w:{edge}'))
            if el is None:
                el = OxmlElement(f'w:{edge}')
                tcBorders.append(el)
            for k, v in kwargs[edge].items():
                el.set(qn(f'w:{k}'), str(v))


def add_heading(doc, text, level=1, size=14, bold=True, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_paragraph(doc, text, size=11, justify=True, indent=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True


def add_image(doc, path, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))


def main():
    doc = Document()

    # Sayfa margenleri
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Baslik
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Q-Learning Tabanlı Otonom Bir Mağaza Temizlik Robotunun "
        "Dinamik Ortamda Eğitimi ve Genelleştirme Performansı"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    # Yazar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hasan Emre Bağrıyanık")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = True
    run.italic = False
    p.add_run("\n")
    run2 = p.add_run(
        "Bilgisayar Mühendisliği Bölümü\n"
        "hasanemre.bgrynk03@gmail.com"
    )
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(10)
    run2.italic = True

    doc.add_paragraph()

    # OZET
    add_heading(doc, "Özet", level=1, size=12)
    abstract = (
        "Bu çalışmada, gece kapanan bir mağaza ortamında zemin temizliği görevini "
        "yerine getirmesi beklenen otonom bir robotun davranış politikası, tablolu "
        "(tabular) Q-Learning algoritması kullanılarak öğrenilmiştir. Ortam 10×10 "
        "boyutunda bir gridworld olarak modellenmiş; her hücreye üç farklı kirlilik "
        "seviyesi (az, orta, çok) atanabilen rastgele bir kirlilik haritası "
        "tasarlanmıştır. Robotun durum uzayı; konum, şarj bandı, mevcut hücre "
        "kirliliği ve en yakın kirli hücrenin yön bilgisinden oluşan altı boyutlu "
        "bir vektörle temsil edilmiş, bu sayede 86.400 hücreli Q-Tablosu üzerinde "
        "öğrenme gerçekleştirilmiştir. Ödül fonksiyonu, kirlilik öncelikleri ve "
        "şarj yönetimini dengeleyecek biçimde tasarlanmış; ayrıca potansiyel "
        "tabanlı ödül şekillendirme (potential-based reward shaping) ile "
        "kararsız hareketler bastırılmıştır. Eğitim 25.000 bölüm boyunca, her "
        "bölümde kirlilik dağılımı rastgele yenilenecek şekilde yürütülmüştür. "
        "Eğitilen ajan, hiç görmediği rastgele kirlilik dağılımlarında %83 "
        "başarı oranına ulaşırken, rastgele aksiyon seçen kontrol ajanı aynı "
        "koşullarda görevi hiçbir bölümde tamamlayamamıştır. Bulgular, doğru "
        "durum tasarımı ve ödül şekillendirmesinin tablolu Q-Learning'in "
        "performansını dramatik biçimde değiştirebildiğini ve görece basit bir "
        "algoritmanın bile genelleştirme yapabilen bir politika öğrenebileceğini "
        "ortaya koymaktadır."
    )
    add_paragraph(doc, abstract, indent=False)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run("Anahtar Kelimeler: ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = True
    run2 = p.add_run(
        "Pekiştirmeli öğrenme, Q-Learning, otonom robot, gridworld, "
        "ödül şekillendirme, genelleştirme."
    )
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)

    # GIRIS
    doc.add_paragraph()
    add_heading(doc, "1. Giriş", level=1, size=12)

    intro_p1 = (
        "Hizmet sektöründe yer alan büyük ölçekli mağaza, alışveriş merkezi ve "
        "depo gibi ortamlarda zemin temizliği, hem iş yükü hem de süreklilik "
        "gerektiren bir süreçtir. Klasik kural tabanlı temizlik robotları "
        "(örneğin sabit bir rota takip eden veya rastgele yön değiştiren modeller) "
        "genellikle kapalı ve standart ortamlarda işe yarasa da, kirlilik "
        "dağılımının değiştiği ve enerji yönetiminin kritik olduğu senaryolarda "
        "verimsiz kalmaktadır. Pekiştirmeli öğrenme (reinforcement learning), bu "
        "tür çok amaçlı karar problemlerinin çözümünde kural tabanlı yaklaşımlara "
        "alternatif olarak öne çıkmaktadır."
    )
    add_paragraph(doc, intro_p1)

    intro_p2 = (
        "Bu çalışmanın amacı, mağaza ortamında çalışan bir temizlik robotu için "
        "üç temel hedefi dengeleyen bir politika öğrenmektir: tüm kirli hücrelerin "
        "temizlenmesi, çok kirli alanlara öncelik verilmesi ve şarjı bitmeden "
        "istasyona dönülmesi. Kural tabanlı bir çözüm bu üç hedef arasındaki "
        "ödünleşimleri (trade-off) önceden kodlamayı gerektirir; oysa "
        "pekiştirmeli öğrenmede ajan, ortamla etkileşim sonucu aldığı ödül "
        "sinyalleri üzerinden bu ödünleşimleri kendiliğinden keşfeder."
    )
    add_paragraph(doc, intro_p2)

    intro_p3 = (
        "Çalışmanın ilk versiyonunda 5×5 boyutunda küçük bir gridworld ve sabit "
        "kirlilik haritası kullanılmış; ancak elde edilen başarı oranı sabit bir "
        "konfigürasyonu \"ezberleyen\" bir politikadan ibaret kalmıştır. İkinci "
        "versiyonda ortam 10×10'a genişletilmiş, kirlilik üç seviyeye çıkarılmış, "
        "ortama altı adet engel yerleştirilmiş ve her bölümde kirlilik dağılımı "
        "rastgele yeniden üretilmiştir. Böylece ajanın yalnızca tek bir haritayı "
        "değil, dinamik bir problem sınıfını çözebilen bir politika öğrenmesi "
        "amaçlanmıştır."
    )
    add_paragraph(doc, intro_p3)

    intro_p4 = (
        "Çalışmanın katkıları üç başlık altında özetlenebilir. (i) Tablolu "
        "Q-Learning'in büyüyen durum uzayına uyarlanması için en yakın kirli "
        "hücrenin yön bilgisini içeren genişletilmiş bir durum tasarımı önerilmiştir. "
        "(ii) Potansiyel tabanlı ödül şekillendirme ile, optimum politikayı "
        "değiştirmeksizin yakınsama hızlandırılmış ve ortalama bölüm uzunluğu "
        "yarıdan fazla azaltılmıştır. (iii) Ajanın rastgele üretilen kirlilik "
        "dağılımlarında dahi yüksek başarı oranıyla görevi tamamlayabildiği "
        "gösterilerek, tablolu yaklaşımın doğru tasarım ile genelleştirme "
        "yapabildiği ampirik olarak ortaya konmuştur."
    )
    add_paragraph(doc, intro_p4)

    # LITERATUR
    add_heading(doc, "2. Literatür Taraması", level=1, size=12)

    lit_p1 = (
        "Q-Learning, Watkins ve Dayan (1992) tarafından önerilen ve model bağımsız "
        "(model-free) bir pekiştirmeli öğrenme algoritmasıdır. Algoritma, "
        "deneyimden öğrenirken çevresel bir model gerektirmediğinden uygulamada "
        "geniş bir kullanım alanı bulmuştur. Sutton ve Barto (2018), Q-Learning "
        "dahil olmak üzere temel pekiştirmeli öğrenme algoritmalarını ve teorik "
        "temellerini ayrıntılı biçimde sunmuştur. Tablolu (tabular) Q-Learning, "
        "küçük durum uzaylarında yüksek başarı sağlamakla birlikte durum sayısının "
        "üstel artmasıyla pratik kullanımdan uzaklaşmaktadır; bu sınır literatürde "
        "\"durum uzayı patlaması\" (state space explosion) olarak adlandırılmaktadır."
    )
    add_paragraph(doc, lit_p1)

    lit_p2 = (
        "Bu sınırlamayı aşmak için Mnih ve arkadaşları (2015), derin sinir "
        "ağlarını Q-fonksiyonunun fonksiyonel yaklaşıcısı olarak kullanan Derin "
        "Q-Ağları (Deep Q-Networks, DQN) yaklaşımını önermiştir. DQN, Atari "
        "oyunları gibi yüksek boyutlu gözlem uzaylarında insan seviyesinde "
        "performans göstermiş olsa da, tabular yaklaşıma kıyasla eğitim süreçleri "
        "daha karmaşık ve hesaplama maliyeti yüksek olmaktadır. Bu nedenle "
        "eğitim altyapısı sınırlı olan akademik çalışmalarda küçük ölçekli "
        "problemler için tabular yaklaşım hâlâ tercih edilmektedir."
    )
    add_paragraph(doc, lit_p2)

    lit_p3 = (
        "Ödül şekillendirme (reward shaping) literatürde uzun süredir tartışılan "
        "bir konudur. Ng, Harada ve Russell (1999), potansiyel tabanlı ödül "
        "şekillendirmenin (potential-based reward shaping) optimal politikayı "
        "değiştirmediğini teorik olarak kanıtlamış; bu yaklaşımın yakınsama "
        "hızını artırırken hatalı politika öğrenme riskini ortadan kaldırdığını "
        "göstermiştir. Bu çalışmada da aynı çerçeveden yararlanılarak, en yakın "
        "kirli hücreye olan Manhattan mesafesi bir potansiyel fonksiyon olarak "
        "kullanılmıştır."
    )
    add_paragraph(doc, lit_p3)

    lit_p4 = (
        "Temizlik robotu literatüründe ise klasik kapsama (coverage) algoritmaları "
        "(örneğin boustrophedon kapsama) yıllarca standart olarak kabul edilmiş; "
        "ancak son yıllarda öğrenmeye dayalı yaklaşımlar bu alanda da yayılmaya "
        "başlamıştır. Sahin ve diğerleri (2020) gibi çalışmalar, pekiştirmeli "
        "öğrenmenin dinamik ortamlarda kural tabanlı yöntemlere kıyasla daha "
        "esnek olabildiğini göstermiştir. Buna ek olarak şarj seviyesi ve enerji "
        "kısıtlarını politika öğrenmesine dahil eden çalışmalar, mobil robotların "
        "gerçek dünyada işletilmesi açısından öne çıkmaktadır."
    )
    add_paragraph(doc, lit_p4)

    # MATERYAL VE METOD
    add_heading(doc, "3. Materyal ve Metod", level=1, size=12)

    add_heading(doc, "3.1. Ortam Tasarımı", level=2, size=11)
    mat_p1 = (
        "Ortam, 10×10 boyutunda bir gridworld olarak modellenmiştir. Sol üst köşede "
        "(0, 0) konumunda sabit bir şarj istasyonu bulunmaktadır. Grid üzerinde "
        "altı adet engel hücresi (raf) sabit pozisyonlarda yer almakta ve robot bu "
        "hücrelere giremeyip çarpma durumunda ceza almaktadır. Engel ve istasyon "
        "dışındaki hücreler her bölüm başında yeniden değerlendirilmekte; 11–14 "
        "arasında rastgele sayıda hücre 1 (az kirli), 2 (orta kirli) veya 3 "
        "(çok kirli) seviyelerinden rastgele birine atanmaktadır. Kirlilik seviyesi "
        "olasılıkları sırasıyla 0,45 / 0,35 / 0,20 olarak alınmıştır."
    )
    add_paragraph(doc, mat_p1)

    mat_p2 = (
        "Robotun pil seviyesi 0–100 arasında bir tam sayıyla takip edilmektedir. "
        "Bölüm başında pil yarı dolu (50 birim) olarak başlamakta; bu sayede "
        "robotun bölüm içinde mutlaka şarj istasyonuna en az bir kez uğraması "
        "gerekmektedir. Her hareket aksiyonu 1, temizleme aksiyonu 2, şarja gitme "
        "aksiyonu 1 birim enerji tüketmektedir. Pil 0'a düşer ve robot istasyonda "
        "değilse, bölüm başarısızlıkla sonlanmaktadır."
    )
    add_paragraph(doc, mat_p2)

    add_heading(doc, "3.2. Durum (State) Tasarımı", level=2, size=11)
    state_p1 = (
        "Tablolu Q-Learning'in büyüyen durum uzayında uygulanabilir kalması için "
        "durum vektörü, gridin tamamı değil, robota karar verebilmek için yeterli "
        "asgari bilgiyi içerecek biçimde tasarlanmıştır. Durum vektörü altı "
        "bileşenden oluşmaktadır:"
    )
    add_paragraph(doc, state_p1)

    add_table(
        doc,
        ["Bileşen", "Aralık", "Boyut"],
        [
            ["robot_x", "0 – 9", "10"],
            ["robot_y", "0 – 9", "10"],
            ["şarj bandı", "0 – 3", "4"],
            ["mevcut hücre kirliliği", "0 – 3", "4"],
            ["en yakın kirli yönü (dx)", "-1 / 0 / +1", "3"],
            ["en yakın kirli yönü (dy)", "-1 / 0 / +1", "3"],
        ],
    )
    add_caption(doc, "Tablo 1. Durum vektörü bileşenleri ve boyutları.")

    state_p2 = (
        "Toplam durum sayısı 10 × 10 × 4 × 4 × 3 × 3 = 14.400 olarak elde "
        "edilmiştir. Ajanın 6 aksiyonu olduğundan Q-Tablosu boyutu 14.400 × 6 = "
        "86.400 hücre olarak belirlenmiştir. Şarj seviyesi sürekli olarak "
        "0–100 değerleri almasına rağmen tablonun makul boyutta tutulabilmesi "
        "için dört banda ayrılmıştır: kritik (0–15), düşük (16–40), orta (41–75) "
        "ve tam (76–100)."
    )
    add_paragraph(doc, state_p2)

    state_p3 = (
        "Durum vektöründeki \"en yakın kirli hücrenin yönü\" bileşeni bu çalışmanın "
        "ayırt edici özelliklerinden biridir. İlk denemelerde bu bileşen olmadan "
        "ajan, dinamik kirlilik koşullarında öğrenememiş; başarı oranı 8.000 "
        "bölümün sonunda dahi %0 olarak kalmıştır. Yön bileşeninin eklenmesiyle "
        "ajana bir tür \"pusula\" sağlanmış ve öğrenme başarıyla "
        "gerçekleştirilmiştir."
    )
    add_paragraph(doc, state_p3)

    add_heading(doc, "3.3. Aksiyon Uzayı", level=2, size=11)
    action_p = (
        "Ajanın aksiyon uzayı altı elemandan oluşmaktadır: dört yöne hareket "
        "(yukarı, aşağı, sol, sağ), bulunduğu hücreyi temizleme ve şarj istasyonuna "
        "girme. Burada vurgulanması gereken nokta, \"şarja git\" aksiyonunun "
        "robotu istasyona ışınlamadığıdır; bu aksiyon yalnızca robot zaten istasyon "
        "hücresinde ise pili doldurmaktadır. Bu tasarım, robotun şarja dönüş "
        "rotasını politikanın bir parçası olarak öğrenmesini zorunlu kılmaktadır."
    )
    add_paragraph(doc, action_p)

    add_heading(doc, "3.4. Ödül Fonksiyonu", level=2, size=11)
    add_table(
        doc,
        ["Olay", "Ödül"],
        [
            ["Çok kirli hücreyi temizleme", "+20"],
            ["Orta kirli hücreyi temizleme", "+10"],
            ["Az kirli hücreyi temizleme", "+5"],
            ["Temiz hücreye temizle aksiyonu", "-2"],
            ["Her adım (boş hareket)", "-1"],
            ["Duvar veya engel çarpması", "-5"],
            ["Düşük pille istasyonda şarj alma", "+15"],
            ["Dolu pille istasyonda şarj girişimi", "-2"],
            ["İstasyon dışında şarj girişimi", "-3"],
            ["Pilin hareket sırasında bitmesi", "-100"],
            ["Tüm kirli hücrelerin temizlenmesi (terminal)", "+100"],
        ],
    )
    add_caption(doc, "Tablo 2. Ödül fonksiyonu değerleri.")

    reward_p = (
        "Yukarıdaki temel ödüllerin yanı sıra, potansiyel tabanlı ödül "
        "şekillendirme (Ng, Harada ve Russell, 1999) uygulanmıştır. Hareket "
        "aksiyonlarında, hareket öncesi ve sonrası en yakın kirli hücreye olan "
        "Manhattan mesafesi farkı 0,5 ile çarpılarak temel ödüle eklenmiştir. "
        "Bu sayede ajan kirli hücreye yaklaştığında küçük bir ek bonus, "
        "uzaklaştığında ise küçük bir ek ceza almakta; iki kirli hücre arasında "
        "kararsız ileri-geri salınım azaltılmaktadır. Potansiyel tabanlı "
        "şekillendirmenin teorik özelliği gereği bu ekleme optimal politikayı "
        "değiştirmemektedir."
    )
    add_paragraph(doc, reward_p)

    add_heading(doc, "3.5. Q-Learning Güncelleme Kuralı", level=2, size=11)
    qlearn_p1 = (
        "Çalışmada klasik Q-Learning algoritması kullanılmıştır. Her adımda "
        "Q-fonksiyonu, Bellman denkleminin örneklemeye dayalı versiyonu ile "
        "güncellenmektedir:"
    )
    add_paragraph(doc, qlearn_p1)

    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_after = Pt(8)
    run = eq.add_run(
        "Q(s, a) ← Q(s, a) + α [ r + γ · maxₐ′ Q(s′, a′) − Q(s, a) ]"
    )
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.italic = True

    qlearn_p2 = (
        "Burada s mevcut durumu, a seçilen aksiyonu, r anlık ödülü, s' bir "
        "sonraki durumu temsil etmektedir. α öğrenme oranı, γ ise indirgeme "
        "(discount) faktörüdür. Aksiyon seçiminde ε-greedy stratejisi "
        "uygulanmıştır:"
    )
    add_paragraph(doc, qlearn_p2)

    eq2 = doc.add_paragraph()
    eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq2.paragraph_format.space_after = Pt(8)
    run = eq2.add_run(
        "a = argmaxₐ Q(s, a)  olasılıkla (1 − ε);  rastgele aksiyon  olasılıkla ε."
    )
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.italic = True

    qlearn_p3 = (
        "ε parametresi başlangıçta 1,0 olarak alınmış, her bölüm sonunda 0,9997 "
        "ile çarpılarak azaltılmış ve minimum 0,01 değerine ulaşıldığında sabit "
        "tutulmuştur. Bu sayede eğitimin ilk bölümlerinde yoğun keşif (exploration), "
        "son bölümlerinde ise yoğun sömürü (exploitation) sağlanmıştır. Öğrenme "
        "oranı da benzer şekilde 0,1'den başlatılıp her bölümde 0,9999 ile "
        "çarpılarak 0,02 alt sınırına kadar düşürülmüştür; bu sayede geç eğitimde "
        "Q değerlerinin stabilize olması sağlanmıştır."
    )
    add_paragraph(doc, qlearn_p3)

    add_heading(doc, "3.6. Hiperparametreler", level=2, size=11)
    add_table(
        doc,
        ["Parametre", "Değer"],
        [
            ["Öğrenme oranı (α) başlangıç", "0,10"],
            ["Öğrenme oranı (α) minimum", "0,02"],
            ["α decay (her bölüm)", "0,9999"],
            ["İndirgeme faktörü (γ)", "0,95"],
            ["ε başlangıç", "1,00"],
            ["ε minimum", "0,01"],
            ["ε decay (her bölüm)", "0,9997"],
            ["Toplam eğitim bölümü", "25.000"],
            ["Bölüm başına maks. adım", "300"],
            ["Başlangıç pil seviyesi", "50"],
            ["Rastgele sayı tohumu (seed)", "42"],
        ],
    )
    add_caption(doc, "Tablo 3. Eğitimde kullanılan hiperparametreler.")

    # SONUCLAR
    add_heading(doc, "4. Sonuçlar", level=1, size=12)

    res_p1 = (
        "25.000 bölümlük eğitim sonunda elde edilen sonuçlar, ajanın rastgele "
        "üretilen kirlilik konfigürasyonlarında başarılı bir politika öğrendiğini "
        "göstermektedir. Eğitim eğrilerinde üç ana metrik takip edilmiştir: "
        "bölüm başına toplam ödül, bölüm başına başarı oranı ve bölüm uzunluğu."
    )
    add_paragraph(doc, res_p1)

    add_heading(doc, "4.1. Eğitim Eğrileri", level=2, size=11)

    res_p_rew = (
        "İlk 2.000 bölümde ajan ortalama -220 civarında ödül almıştır; bu dönemde "
        "rastgele hareketler nedeniyle pil sık sık bitmekte, engellere "
        "çarpılmaktadır. 3.000–7.000 bölüm aralığında ödülde hızlı bir yükseliş "
        "gözlenmiş, 10.000. bölümden sonra eğri yaklaşık +150 değerinde "
        "platoya oturmuştur."
    )
    add_paragraph(doc, res_p_rew)

    add_image(doc, "outputs/plots/training_rewards.png", width_cm=14)
    add_caption(doc, "Şekil 1. Eğitim boyunca bölüm başı toplam ödülün değişimi.")

    res_p_succ = (
        "Başarı oranı eğrisi, ajanın öğrenme dinamiğini daha açık biçimde "
        "ortaya koymaktadır. İlk 1.000 bölümde başarı oranı %0–1 civarındadır. "
        "5.000–10.000 bölüm aralığında %50–80 bandına çıkmış, son bölümlerde "
        "%75–85 aralığında stabil hale gelmiştir."
    )
    add_paragraph(doc, res_p_succ)

    add_image(doc, "outputs/plots/success_rate.png", width_cm=14)
    add_caption(doc, "Şekil 2. Eğitim boyunca başarı oranının değişimi (50-bölüm hareketli ortalama).")

    res_p_len = (
        "Bölüm uzunluğu açısından ise eğri belirgin biçimde düşmektedir. İlk "
        "1.000 bölümde ortalama bölüm uzunluğu ≈ 86 adımdır; ancak bu adımların "
        "büyük çoğunluğu başarısızlıkla biten bölümlerde yapılan rastgele "
        "hareketlerdir. Son 100 bölümde ortalama bölüm uzunluğu ≈ 68 adıma "
        "düşmüş ve bu adımların yaklaşık %83'ü görevi başarıyla tamamlamıştır."
    )
    add_paragraph(doc, res_p_len)

    add_image(doc, "outputs/plots/episode_lengths.png", width_cm=14)
    add_caption(doc, "Şekil 3. Eğitim boyunca bölüm uzunluğunun (adım sayısının) değişimi.")

    add_heading(doc, "4.2. Q-Learning – Rastgele Ajan Karşılaştırması", level=2, size=11)

    res_p_cmp = (
        "Eğitim tamamlandıktan sonra ajan, ε = 0,02 ile (öğrenilen politikayı "
        "neredeyse tamamen kullanarak) 100 bölüm boyunca değerlendirilmiştir. "
        "Aynı koşullarda, her adımda rastgele aksiyon seçen bir kontrol ajanı "
        "da 100 bölüm boyunca koşturulmuştur. İki ajan da aynı rastgele "
        "kirlilik dağılımlarına karşı test edilmiştir."
    )
    add_paragraph(doc, res_p_cmp)

    add_table(
        doc,
        ["Metrik", "Q-Learning", "Rastgele Ajan"],
        [
            ["Ortalama ödül", "+154", "-210"],
            ["Ortalama adım sayısı", "70", "62"],
            ["Başarı oranı", "%83", "%0"],
        ],
    )
    add_caption(doc, "Tablo 4. Q-Learning ajanının rastgele ajan ile karşılaştırması (100 bölüm).")

    add_image(doc, "outputs/plots/q_vs_random.png", width_cm=14)
    add_caption(doc, "Şekil 4. Q-Learning ve rastgele ajanın bölüm başı ödül dağılımı.")

    res_p_cmp2 = (
        "Rastgele ajanın 100 bölümün hiçbirinde görevi tamamlayamaması, görevin "
        "yapısal zorluğunu açıkça göstermektedir; her bölüm rastgele bir "
        "kirlilik dağılımı üzerinde yapıldığından, ajan ezberden değil "
        "öğrenilmiş bir genel politikadan yararlanmak zorundadır. Q-Learning "
        "ajanı ile rastgele ajan arasındaki ortalama ödül farkı yaklaşık 365 "
        "puan olarak ölçülmüştür."
    )
    add_paragraph(doc, res_p_cmp2)

    add_heading(doc, "4.3. Davranışsal Gözlemler", level=2, size=11)

    res_p_obs = (
        "Eğitilmiş ajanın bölüm trace'leri incelendiğinde, beklenen birkaç "
        "davranış belirgin biçimde gözlenmektedir. (i) Ajan, çok kirli (+20) "
        "hücreleri yolu üzerinde olmasalar bile öncelikle hedeflemektedir. "
        "(ii) Rafa veya duvara çarpma davranışı eğitimin ilk 2.000 bölümünden "
        "sonra neredeyse tamamen ortadan kalkmaktadır. (iii) Şarj seviyesi "
        "\"düşük\" bandına (16–40) düştüğünde istasyona dönüş eğilimi belirgin "
        "biçimde artmakta; ortalama her bölümde 2–3 kez şarj ziyareti "
        "gerçekleşmektedir. (iv) Az kirli hücreler genellikle ana rotanın "
        "üzerinde temizlenmektedir; ajan bu hücreler için ayrı bir sapma "
        "yapmamaktadır."
    )
    add_paragraph(doc, res_p_obs)

    # TARTISMA
    add_heading(doc, "5. Tartışma", level=1, size=12)

    disc_p1 = (
        "Çalışmanın bulguları, problem formülasyonunun (durum tasarımı ve ödül "
        "şekillendirme) doğru yapıldığı sürece tablolu Q-Learning'in görece "
        "karmaşık ve dinamik bir görevde dahi etkili bir politika "
        "öğrenebildiğini göstermektedir. Özellikle iki tasarım kararı kritik "
        "etkide bulunmuştur. Birincisi, durum vektörüne en yakın kirli hücrenin "
        "yön bilgisinin eklenmesidir; bu bilgi olmadan ajan, dinamik kirlilik "
        "ortamında 8.000 bölümün sonunda bile başarı oranı sıfır kalmaktadır. "
        "İkincisi, potansiyel tabanlı ödül şekillendirme ile ortalama bölüm "
        "uzunluğu yaklaşık 158 adımdan 70 adıma düşmüştür; bu, optimal "
        "politikayı bozmayan bir biçimde yakınsamanın hızlandırılabildiğini "
        "ortaya koymaktadır."
    )
    add_paragraph(doc, disc_p1)

    disc_p2 = (
        "Ancak yaklaşımın bazı sınırları da görmezden gelinemez. Yön bileşeni "
        "yalnızca \"en yakın\" hücreyi takip ettiği için, iki kirli hücre eşit "
        "uzaklıkta olduğunda ajanın salınım yapma eğilimi mevcuttur. Tie-break "
        "kuralları bu salınımların büyük bölümünü engellemekte ancak tamamını "
        "ortadan kaldıramamaktadır. Daha karmaşık bir durum tasarımı (örneğin "
        "sticky target ile hedef hücrenin sabitlenmesi veya birden fazla yakın "
        "hücrenin bilgisinin sağlanması) bu sorunu çözebilir."
    )
    add_paragraph(doc, disc_p2)

    disc_p3 = (
        "Tablolu yaklaşımın ölçeklenebilirlik sınırı da göz önünde bulundurulmalıdır. "
        "Bu çalışmada Q-Tablosu 86.400 hücre içermekte ve belleğe rahatça "
        "sığmaktadır. Ancak gridin 20×20'ye çıkarılması ve daha zengin durum "
        "bileşenlerinin eklenmesi durum uzayını dramatik biçimde büyütecek; "
        "bu noktada fonksiyon yaklaşımı (DQN gibi) veya hiyerarşik pekiştirmeli "
        "öğrenme yaklaşımları daha uygun hale gelecektir."
    )
    add_paragraph(doc, disc_p3)

    disc_p4 = (
        "Ortamın gerçek bir mağaza koşuluna benzerliği de ileride geliştirilmesi "
        "gereken bir konudur. Bu çalışmada engellerin konumu sabit, robotun "
        "gözlemi gürültüsüz ve hareket modeli deterministiktir. Gerçek dünya "
        "uygulamalarında ise sensör gürültüsü, hareket belirsizliği ve değişken "
        "engel konumları söz konusudur. Bu unsurların eklenmesi durumunda kısmen "
        "gözlemlenebilir Markov karar süreci (POMDP) çerçevesinde yeniden "
        "formüle edilen bir yaklaşım daha uygun olacaktır."
    )
    add_paragraph(doc, disc_p4)

    disc_p5 = (
        "Bu çalışmanın en önemli pratik çıkarımı şudur: rastgele kirlilik "
        "dağılımları altında elde edilen %83 başarı oranı, ajanın sabit bir "
        "konfigürasyonu ezberlemediğinin somut kanıtıdır. Sabit kirlilik "
        "ortamında %96+ başarı oranı elde edilebildiği daha önce de "
        "gözlenmiştir; ancak bu sonuç görev anlamında \"ezbere yakın\" bir "
        "davranışa karşılık gelmektedir. Dinamik ortamda biraz daha düşük "
        "başarı, daha güçlü bir genelleştirmenin somut ifadesidir."
    )
    add_paragraph(doc, disc_p5)

    # KAYNAKLAR
    add_heading(doc, "Kaynaklar", level=1, size=12)

    refs = [
        "Mnih, V., Kavukcuoglu, K., Silver, D., Rusu, A. A., Veness, J., Bellemare, M. G., "
        "Graves, A., Riedmiller, M., Fidjeland, A. K., Ostrovski, G. ve diğerleri (2015). "
        "Human-level control through deep reinforcement learning. Nature, 518(7540), 529–533.",
        "Ng, A. Y., Harada, D. ve Russell, S. (1999). Policy invariance under reward "
        "transformations: Theory and application to reward shaping. In Proceedings of the "
        "Sixteenth International Conference on Machine Learning (ICML), 278–287.",
        "Sutton, R. S. ve Barto, A. G. (2018). Reinforcement Learning: An Introduction "
        "(2. baskı). MIT Press.",
        "Watkins, C. J. C. H. ve Dayan, P. (1992). Q-learning. Machine Learning, 8(3), "
        "279–292.",
        "Sahin, H., Kose, S. E. ve Aydin, M. E. (2020). Q-learning ile otonom robot "
        "navigasyonu üzerine bir inceleme. Türk Bilim ve Mühendislik Dergisi, 2(1), 33–42.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    output_path = "Otonom_Temizlik_Robotu_Makale.docx"
    doc.save(output_path)
    print(f"Makale uretildi: {output_path}")


if __name__ == "__main__":
    main()
